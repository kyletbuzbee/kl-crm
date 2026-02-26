"""
Deeper inspection: show ALL cell content including multiline text,
check actual column count, and find where Net Weight lives.
"""
from docx import Document

doc = Document('transactions.docx')

table = doc.tables[0]
print(f"Table columns: {len(table.columns)}")
print(f"Table rows: {len(table.rows)}")
print()

# Show FIRST 15 rows raw - including any newlines in cells
print("=== FIRST 15 ROWS (raw) ===")
for r_idx, row in enumerate(table.rows[:15]):
    print(f"\n--- Row {r_idx} ---")
    for c_idx, cell in enumerate(row.cells):
        raw = repr(cell.text)
        print(f"  col[{c_idx}]: {raw}")

# Now scan for rows where col[5] is NOT empty
print("\n\n=== ROWS WHERE COL[5] IS NOT EMPTY (first 20) ===")
count = 0
for r_idx, row in enumerate(table.rows):
    cells = [cell.text for cell in row.cells]
    if len(cells) > 5 and cells[5].strip():
        print(f"  Row {r_idx}: col[5]={repr(cells[5])!r}  full={[repr(c) for c in cells]}")
        count += 1
        if count >= 20:
            break

if count == 0:
    print("  (none found - col[5] is always empty)")

# Also check if any cell contains a number followed by another number (weight + price together)
print("\n\n=== ROWS WHERE ANY CELL CONTAINS \\n (newline in cell) ===")
count = 0
for r_idx, row in enumerate(table.rows[:50]):
    for c_idx, cell in enumerate(row.cells):
        if '\n' in cell.text:
            print(f"  Row {r_idx} col[{c_idx}]: {repr(cell.text)}")
            count += 1
if count == 0:
    print("  (no multiline cells found in first 50 rows)")

# Show paragraphs inside each cell for row 1 (likely data row)
print("\n\n=== PARAGRAPH BREAKDOWN FOR ROWS 1-5 ===")
for r_idx, row in enumerate(table.rows[1:6], start=1):
    print(f"\n  Row {r_idx}:")
    for c_idx, cell in enumerate(row.cells):
        paras = [repr(p.text) for p in cell.paragraphs if p.text.strip()]
        if paras:
            print(f"    col[{c_idx}] paragraphs: {paras}")
