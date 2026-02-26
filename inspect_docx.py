from docx import Document

doc = Document('transactions.docx')

print('=== PARAGRAPHS (non-empty) ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] {repr(p.text)}')

print()
print('=== TABLES ===')
for t_idx, table in enumerate(doc.tables):
    print(f'Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols')
    for r_idx, row in enumerate(table.rows):
        cells = [repr(cell.text.strip()) for cell in row.cells]
        print(f'  Row {r_idx}: {cells}')
