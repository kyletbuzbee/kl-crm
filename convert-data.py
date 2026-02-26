from docx import Document
import csv
import os
import re

# Paths
docx_path = r"D:\K&L Recycling Outreach\Untitled document.docx"
csv_path = r"D:\K&L Recycling Outreach\INFORMATION_FOR_KYLE.csv"

def final_attempt_scan():
    if not os.path.exists(docx_path):
        print(f"Error: Could not find {docx_path}")
        return

    print("Running Advanced Scan on K&L Report...")
    doc = Document(docx_path)
    all_data = []

    # Pattern: 4-6 digit ticket followed by XX/XX/XXXX date
    # This pattern is the most consistent anchor in your report
    pattern = re.compile(r'(\d{4,6})\s+(\d{2}/\d{2}/\d{4})')

    # Scan paragraphs AND tables (to cover all bases)
    search_targets = doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            search_targets.extend(row.cells)

    for item in search_targets:
        text = " ".join(item.text.split())
        if not text:
            continue
            
        match = pattern.search(text)
        if match:
            # We found a transaction line
            parts = text.split()
            
            # K&L Report logic: 
            # Ticket=parts[0], Date=parts[1], Price=Last, Net=Second to Last
            if len(parts) >= 5:
                ticket = parts[0]
                date = parts[1]
                # Price often has a negative sign for fees, so we grab it carefully
                price = parts[-1]
                net = parts[-2]
                details = " ".join(parts[2:-2])
                
                all_data.append([ticket, date, details, net, price])

    if all_data:
        # Sort and remove duplicates (common when scanning both paragraphs and tables)
        unique_data = list({tuple(row): row for row in all_data}.values())
        
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Ticket #", "Date", "Supplier/Material", "Net", "Price"])
            writer.writerows(unique_data)
        print(f"Success! Captured {len(unique_data)} transactions.")
        print(f"CSV saved to: {csv_path}")
    else:
        print("No transactions found. Verify the Word document contains text and not just images.")

if __name__ == "__main__":
    final_attempt_scan()